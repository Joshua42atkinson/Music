#!/usr/bin/env python3
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def apply_apa_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 2.0
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)

def create_title_page(doc):
    for _ in range(4):
        doc.add_paragraph()
        
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Voix Vive Maturation Map: A Pedagogical Architecture")
    run.bold = True
    
    doc.add_paragraph()
    
    details = [
        "Joshua Atkinson",
        "Learning Design and Technology Program",
        "Department of Curriculum and Instruction, Purdue University",
        "EDCI 57300: LEARNING DESIGN AND TECHNOLOGY and CORPORATE TRAINING AND COMMUNICATION PROGRAM PRACTICUM",
        "Instructor: Dr. Jennifer C. Richardson",
        "jennrich@purdue.edu",
        "June 2026"
    ]
    
    for text in details:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_page_break()

def parse_inline_markdown(p, text):
    """Basic parser for **bold** and *italic* and `code`."""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('*') and part.endswith('*'):
            r = p.add_run(part[1:-1])
            r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1])
            r.font.name = 'Courier New'
        else:
            p.add_run(part)

def set_table_borders(table):
    """
    Apply APA-like borders to a table:
    Top border, bottom border of header, and bottom border of the whole table.
    No vertical borders.
    """
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'bottom', 'insideH']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
        
    tblPr = tbl.tblPr
    tblPr.append(tblBorders)

def process_table(doc, table_lines):
    if not table_lines:
        return
        
    # Remove alignment row (e.g. |---|---|)
    cleaned_lines = [line for line in table_lines if '---' not in line]
    if not cleaned_lines:
        return
        
    # Parse rows
    parsed_rows = []
    for line in cleaned_lines:
        # Strip outer pipes
        if line.startswith('|'): line = line[1:]
        if line.endswith('|'): line = line[:-1]
        
        cells = [cell.strip() for cell in line.split('|')]
        parsed_rows.append(cells)
        
    if not parsed_rows:
        return
        
    # Add table title (APA format)
    p = doc.add_paragraph()
    r = p.add_run("Table 1\n")
    r.bold = True
    r2 = p.add_run("Pedagogical Claims vs. System Functions")
    r2.italic = True
    p.paragraph_format.line_spacing = 2.0
        
    num_cols = len(parsed_rows[0])
    table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
    
    for i, row in enumerate(parsed_rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                p = cell.paragraphs[0]
                parse_inline_markdown(p, cell_text)
                
    set_table_borders(table)
    doc.add_paragraph() # Add space after table

def generate_apa_doc():
    doc = Document()
    apply_apa_style(doc)
    create_title_page(doc)
    
    md_path = 'src/assets/MaturationMap.md'
    if not os.path.exists(md_path):
        print(f"File {md_path} not found.")
        return
        
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
        
    lines = content.split('\n')
    
    in_code_block = False
    in_table = False
    table_lines = []
    
    for line in lines:
        line_stripped = line.strip()
        
        if line_stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = 'Courier New'
            p.paragraph_format.left_indent = Inches(0.5)
            continue
            
        # Table detection
        if line_stripped.startswith('|'):
            in_table = True
            table_lines.append(line_stripped)
            continue
        elif in_table:
            # Table ended
            process_table(doc, table_lines)
            in_table = False
            table_lines = []
            
        if not line_stripped:
            continue
            
        if line_stripped.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line_stripped[2:].strip())
            run.bold = True
        elif line_stripped.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line_stripped[3:].strip())
            run.bold = True
        elif line_stripped.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line_stripped[4:].strip())
            run.bold = True
        elif line_stripped.startswith('#### '):
            p = doc.add_paragraph()
            run = p.add_run(line_stripped[5:].strip())
            run.bold = True
        elif line_stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_markdown(p, line_stripped[2:])
        elif re.match(r'^\d+\.\s', line_stripped):
            p = doc.add_paragraph(style='List Number')
            # Extract text after "1. "
            text = re.sub(r'^\d+\.\s+', '', line_stripped)
            parse_inline_markdown(p, text)
        elif line_stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            parse_inline_markdown(p, line_stripped[2:])
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            parse_inline_markdown(p, line_stripped)
            
    if in_table:
        process_table(doc, table_lines)

    output_file = 'public/Maturation_Map_Submission.docx'
    doc.save(output_file)
    print(f"Successfully created {output_file}")

if __name__ == "__main__":
    generate_apa_doc()
