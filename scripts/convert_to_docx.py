#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Read the markdown file
with open('EDCI_57300_Project_Proposal.md', 'r') as f:
    content = f.read()

# Create a new Document
doc = Document()

# Split content by lines
lines = content.split('\n')

for line in lines:
    if line.startswith('# '):
        # Main heading
        doc.add_heading(line[2:], level=0)
    elif line.startswith('## '):
        # Section heading
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        # Subsection heading
        doc.add_heading(line[4:], level=2)
    elif line.startswith('#### '):
        # Sub-subsection heading
        doc.add_heading(line[5:], level=3)
    elif line.startswith('**') and line.endswith('**'):
        # Bold text
        p = doc.add_paragraph(line[2:-2])
        p.runs[0].bold = True
    elif line.startswith('- **'):
        # Bold list item
        p = doc.add_paragraph(line[3:], style='List Bullet')
        p.runs[0].bold = True
    elif line.startswith('- '):
        # List item
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('|'):
        # Table row - skip for now (complex parsing)
        continue
    elif line.strip() == '':
        # Empty line
        doc.add_paragraph()
    elif line.startswith('---'):
        # Horizontal rule
        doc.add_paragraph('_' * 50)
    else:
        # Regular paragraph
        if line.strip():
            doc.add_paragraph(line)

# Save the document
doc.save('EDCI_57300_Project_Proposal.docx')
print("Document saved as EDCI_57300_Project_Proposal.docx")
